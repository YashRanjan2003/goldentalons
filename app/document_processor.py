import os
import sys
import pytesseract
import cv2
import spacy
from PIL import Image
import pandas as pd
from pdf2image import convert_from_path
import tempfile
import docx
import re

class DocumentProcessor:
    def __init__(self):
        # Load spaCy model for financial entity recognition
        self.nlp = spacy.load("en_core_web_sm")
        # Define financial entity patterns
        self.financial_patterns = [
            {"label": "AMOUNT", "pattern": [{"LIKE_NUM": True}, {"LOWER": {"IN": ["million", "billion", "trillion"]}, "OP": "?"}]},
            {"label": "DATE", "pattern": [{"IS_DIGIT": True, "LENGTH": 1, "OP": "?"}, {"IS_DIGIT": True, "LENGTH": {">=": 1, "<=": 2}}, 
                                        {"LOWER": {"IN": ["/", "-"]}}, {"IS_DIGIT": True, "LENGTH": {">=": 1, "<=": 2}},
                                        {"LOWER": {"IN": ["/", "-"]}, "OP": "?"}, {"IS_DIGIT": True, "LENGTH": {">=": 2, "<=": 4}, "OP": "?"}]},
            {"label": "PERCENTAGE", "pattern": [{"LIKE_NUM": True}, {"ORTH": "%"}]},
            {"label": "RATE", "pattern": [{"LOWER": "rate"}]},
            {"label": "CURRENCY", "pattern": [{"ORTH": {"IN": ["$", "€", "£", "¥"]}}, {"LIKE_NUM": True}]},
            {"label": "TERM", "pattern": [{"LOWER": {"IN": ["term", "maturity", "duration"]}}, {"LOWER": {"IN": ["sheet", "date", "period"]}, "OP": "?"}]},
            {"label": "ISSUER", "pattern": [{"LOWER": "issuer"}]},
            {"label": "INVESTOR", "pattern": [{"LOWER": {"IN": ["investor", "holder", "buyer"]}}]},
            {"label": "SECURITY", "pattern": [{"LOWER": {"IN": ["bond", "note", "security", "stock", "share", "warrant", "option"]}}]}
        ]
        
        # Add patterns to spaCy
        ruler = self.nlp.add_pipe("entity_ruler", before="ner")
        ruler.add_patterns(self.financial_patterns)

    def process_file(self, file_path):
        """
        Process a file and extract text, tables, and entities based on file type
        """
        if not os.path.exists(file_path):
            return {"error": f"File not found: {file_path}"}

        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
            return self.process_image(file_path)
        elif file_extension == '.pdf':
            return self.process_pdf(file_path)
        elif file_extension == '.docx':
            return self.process_docx(file_path)
        elif file_extension == '.txt':
            return self.process_text_file(file_path)
        else:
            return {"error": f"Unsupported file type: {file_extension}"}

    def process_image(self, image_path):
        """
        Process an image file using OCR
        """
        try:
            # Preprocess image for better OCR
            img = cv2.imread(image_path)
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
            
            # Use Tesseract for OCR
            text = pytesseract.image_to_string(thresh)
            
            # Process the extracted text with spaCy
            entities = self.extract_financial_entities(text)
            
            result = {
                "file_type": "image",
                "extracted_text": text,
                "entities": entities,
                "tables": []  # No tables from images
            }
            
            return result
        except Exception as e:
            return {"error": f"Error processing image: {str(e)}"}

    def process_pdf(self, pdf_path):
        """
        Process a PDF file for text, tables, and entities
        """
        try:
            result = {
                "file_type": "pdf",
                "extracted_text": "",
                "entities": [],
                "tables": []
            }
            
            # Convert PDF to images and extract text with OCR
            all_text = ""
            with tempfile.TemporaryDirectory() as temp_dir:
                images = convert_from_path(pdf_path)
                for i, image in enumerate(images):
                    image_path = os.path.join(temp_dir, f'page_{i}.png')
                    image.save(image_path, 'PNG')
                    
                    # Preprocess image for better OCR
                    img = cv2.imread(image_path)
                    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
                    
                    # Extract text with OCR
                    page_text = pytesseract.image_to_string(thresh)
                    all_text += f"\n\n--- PAGE {i+1} ---\n\n{page_text}"
            
            result["extracted_text"] = all_text
            
            # Simple table detection using regex
            table_lines = []
            current_table = []
            in_table = False
            
            for line in all_text.split('\n'):
                # Look for lines with multiple pipe characters or similar table indicators
                if re.search(r'[\|\+\-]{3,}', line) or line.count('\t') >= 3:
                    if not in_table:
                        in_table = True
                    current_table.append(line)
                elif in_table and line.strip() == '':
                    # End of table detected
                    if len(current_table) > 2:  # Ensure it's actually a table with some rows
                        table_lines.append(current_table)
                    current_table = []
                    in_table = False
                elif in_table:
                    current_table.append(line)
            
            # Add any remaining table
            if in_table and len(current_table) > 2:
                table_lines.append(current_table)
            
            # Convert detected tables to structured format
            for i, table in enumerate(table_lines):
                result["tables"].append({
                    "table_id": i,
                    "page": "unknown",
                    "data": table
                })
            
            # Extract financial entities
            result["entities"] = self.extract_financial_entities(all_text)
            
            return result
        except Exception as e:
            return {"error": f"Error processing PDF: {str(e)}"}

    def process_docx(self, docx_path):
        """
        Process a Word document
        """
        try:
            doc = docx.Document(docx_path)
            full_text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            
            # Extract text from tables
            tables_data = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append({
                    "table_id": i,
                    "data": table_data
                })
            
            # Extract financial entities
            entities = self.extract_financial_entities(full_text)
            
            result = {
                "file_type": "docx",
                "extracted_text": full_text,
                "entities": entities,
                "tables": tables_data
            }
            
            return result
        except Exception as e:
            return {"error": f"Error processing DOCX: {str(e)}"}

    def process_text_file(self, text_path):
        """
        Process a text file
        """
        try:
            with open(text_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Extract financial entities
            entities = self.extract_financial_entities(text)
            
            result = {
                "file_type": "text",
                "extracted_text": text,
                "entities": entities,
                "tables": []  # No tables from text files
            }
            
            return result
        except Exception as e:
            return {"error": f"Error processing text file: {str(e)}"}

    def extract_financial_entities(self, text):
        """
        Extract financial entities from text using spaCy
        """
        doc = self.nlp(text)
        entities = []
        
        # Extract named entities
        for ent in doc.ents:
            entities.append({
                "text": ent.text,
                "label": ent.label_,
                "start": ent.start_char,
                "end": ent.end_char
            })
        
        # Extract custom financial patterns
        amount_pattern = r'\$\s*\d+(?:,\d+)*(?:\.\d+)?(?:\s*(?:million|billion|trillion))?'
        percentage_pattern = r'\d+(?:\.\d+)?\s*%'
        date_pattern = r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2}'
        
        # Find amounts
        for match in re.finditer(amount_pattern, text):
            entities.append({
                "text": match.group(),
                "label": "MONEY",
                "start": match.start(),
                "end": match.end()
            })
        
        # Find percentages
        for match in re.finditer(percentage_pattern, text):
            entities.append({
                "text": match.group(),
                "label": "PERCENTAGE",
                "start": match.start(),
                "end": match.end()
            })
        
        # Find dates
        for match in re.finditer(date_pattern, text):
            entities.append({
                "text": match.group(),
                "label": "DATE",
                "start": match.start(),
                "end": match.end()
            })
        
        return entities

def main():
    if len(sys.argv) < 2:
        print("Usage: python document_processor.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    processor = DocumentProcessor()
    result = processor.process_file(file_path)
    
    # Print results
    if "error" in result:
        print(f"Error: {result['error']}")
    else:
        print(f"\n--- FILE TYPE: {result['file_type'].upper()} ---")
        
        print("\n--- EXTRACTED TEXT (first 500 chars) ---")
        print(result['extracted_text'][:500] + "..." if len(result['extracted_text']) > 500 else result['extracted_text'])
        
        print("\n--- FINANCIAL ENTITIES ---")
        for entity in result['entities']:
            print(f"{entity['label']}: {entity['text']}")
        
        print("\n--- TABLES ---")
        if len(result['tables']) == 0:
            print("No tables found")
        else:
            print(f"Found {len(result['tables'])} tables")
            for table in result['tables'][:2]:  # Show just first 2 tables
                print(f"\nTable {table['table_id'] + 1}:")
                if isinstance(table['data'], list) and all(isinstance(x, list) for x in table['data']):
                    # Print formatted table for docx tables
                    for row in table['data'][:5]:
                        print(" | ".join(row))
                else:
                    # Print raw table data for detected tables in PDFs
                    print("\n".join(table['data'][:10]))
                
                if len(result['tables']) > 2:
                    print("... more tables available ...")

if __name__ == "__main__":
    main() 